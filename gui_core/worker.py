from pathlib import Path
import threading

from PyQt5.QtCore import QThread, pyqtSignal
from docx import Document

from src.inference import translate_many, is_model_loaded


class TranslateWorker(QThread):
    """
    Worker поток:
    - не блокирует UI;
    - переводит DOCX после заранее загруженной модели;
    - сохраняет структуру документа: обычные абзацы и таблицы;
    - НЕ сохраняет файл сам: после перевода UI спросит пользователя, куда сохранить результат;
    - поддерживает мягкую отмену между батчами генерации.
    """

    document_ready = pyqtSignal(object, str)
    cancelled = pyqtSignal(str)
    error = pyqtSignal(str)
    progress = pyqtSignal(int)
    status = pyqtSignal(str)

    def __init__(self, file_path: str):
        super().__init__()
        self.file_path = file_path
        self._cancel_event = threading.Event()

    def request_cancel(self) -> None:
        self._cancel_event.set()

    def is_cancelled(self) -> bool:
        return self._cancel_event.is_set()

    def _raise_if_cancelled(self) -> None:
        if self.is_cancelled():
            raise InterruptedError("Перевод отменён пользователем.")

    def _iter_paragraphs(self, container):
        for paragraph in container.paragraphs:
            yield paragraph

        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from self._iter_paragraphs(cell)

    def read_docx(self):
        doc = Document(self.file_path)
        items = []

        for paragraph in self._iter_paragraphs(doc):
            self._raise_if_cancelled()
            text = paragraph.text.strip()
            if text:
                items.append((paragraph, text))

        return doc, items

    def _replace_paragraph_text(self, paragraph, text: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)

    def _suggested_output_path(self) -> str:
        source = Path(self.file_path)
        return str(source.with_name(source.stem + "_translated.docx"))

    def run(self):
        try:
            if not is_model_loaded():
                raise RuntimeError(
                    "Модель ещё не загружена. Дождитесь сообщения 'Модель готова' и запустите перевод снова."
                )

            self._raise_if_cancelled()
            self.status.emit("Читаю DOCX и ищу текст в абзацах и таблицах...")
            self.progress.emit(2)
            doc, items = self.read_docx()

            self._raise_if_cancelled()
            if not items:
                raise ValueError("В DOCX не найден текст в обычных абзацах или таблицах.")

            texts = [text for _, text in items]
            total = len(texts)
            self.status.emit(f"Найдено текстовых фрагментов: {total}")
            self.progress.emit(5)

            def on_translation_progress(done: int, all_count: int):
                if all_count <= 0:
                    self.progress.emit(5)
                    return
                self.progress.emit(5 + int(done / all_count * 90))

            translated_texts = translate_many(
                texts,
                progress_callback=on_translation_progress,
                status_callback=self.status.emit,
                cancel_callback=self.is_cancelled,
            )

            self._raise_if_cancelled()
            self.status.emit("Записываю перевод в структуру DOCX...")
            for (paragraph, _), translated in zip(items, translated_texts):
                self._raise_if_cancelled()
                self._replace_paragraph_text(paragraph, translated)

            self.progress.emit(97)
            self.status.emit("Перевод готов. Выберите, куда сохранить файл.")
            self.document_ready.emit(doc, self._suggested_output_path())

        except InterruptedError as exc:
            self.cancelled.emit(str(exc))
        except Exception as exc:
            self.error.emit(str(exc))
